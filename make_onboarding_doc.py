"""
make_onboarding_doc.py — Generate the Incoming Science user onboarding Word document.

Run once:
    python make_onboarding_doc.py
Output: docs/incoming_science_onboarding.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

# ── Colour palette (matches digest) ──────────────────────────────────────────
SAGE        = RGBColor(0x8F, 0xAF, 0x8F)   # section headers
DARK        = RGBColor(0x2C, 0x28, 0x26)   # body text
WARM_BG     = RGBColor(0xF7, 0xF4, 0xEF)   # fill for answer boxes
MUTED       = RGBColor(0x5C, 0x55, 0x50)   # hint text


def set_cell_bg(cell, hex_str: str):
    """Apply background shading to a table cell. hex_str e.g. 'F7F4EF'"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_str)
    tcPr.append(shd)


def add_section_header(doc, text):
    """Add a sage-coloured section header paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = SAGE
    run.font.name = "Calibri"
    # Bottom border under header
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "8FAF8F")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_hint(doc, text):
    """Add a small muted hint paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.size  = Pt(9)
        run.font.color.rgb = MUTED
        run.font.italic = True
        run.font.name = "Calibri"
    return p


def add_answer_box(doc, lines=3):
    """Add a shaded single-cell table as an answer box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, "F7F4EF")
    cell.width = Inches(6)
    # Add blank lines inside
    for _ in range(lines):
        p = cell.add_paragraph("")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
    return table


def build():
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Incoming Science")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = DARK
    run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("New User Profile Form")
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED
    run.font.name = "Calibri"
    run.italic = True

    # ── Intro paragraph ───────────────────────────────────────────────────────
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    run = intro.add_run(
        "Incoming Science is a personal arXiv digest tool for researchers. Every morning, "
        "it fetches the latest papers in your field, ranks them by relevance to your research "
        "interests using AI, and delivers a scored PDF to your inbox — ready to read on your phone. "
        "Papers are rated with one tap, and those ratings feed back into an evolving taste profile "
        "that sharpens recommendations over time."
    )
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    run.font.name = "Calibri"

    inst = doc.add_paragraph()
    inst.paragraph_format.space_after = Pt(4)
    run = inst.add_run(
        "Please fill in the four sections below and return this form. "
        "Your answers will be used to build your personal research profile."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED
    run.font.italic = True
    run.font.name = "Calibri"

    # ── Part 1: arXiv categories ──────────────────────────────────────────────
    add_section_header(doc, "Part 1 of 4 — arXiv Categories")
    add_hint(doc, "Which arXiv listing pages are you interested in? Enter one or more categories, "
                  "comma-separated.\n"
                  "Examples:  cond-mat  |  cond-mat.str-el  |  quant-ph  |  cs.AI  |  hep-th")
    add_answer_box(doc, lines=2)

    # ── Part 2: Research interests ────────────────────────────────────────────
    add_section_header(doc, "Part 2 of 4 — Research Interests")
    add_hint(doc, "Describe your research interests in your own words. Mention what you focus on "
                  "most and what you follow more loosely. Be as specific as you can — mention "
                  "techniques, materials, or phenomena you care about.\n"
                  "Aim for 1–2 paragraphs (roughly 100–200 words).")
    add_answer_box(doc, lines=10)

    # ── Part 3: Researchers to follow ────────────────────────────────────────
    add_section_header(doc, "Part 3 of 4 — Researchers to Follow")
    add_hint(doc, "List researchers whose new papers you always want to see. "
                  "Comma-separated or one per line.")
    add_answer_box(doc, lines=5)

    # ── Part 4: Recently read papers ─────────────────────────────────────────
    add_section_header(doc, "Part 4 of 4 — Recently Read Papers")
    add_hint(doc, "List papers you have read recently that represent your interests well. "
                  "Accepted formats: arXiv URLs, arXiv IDs (e.g. 2301.12345), DOI links, "
                  "or any journal page URL. Add as many rows as you like.")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for cell, text in zip(hdr_cells, ["Paper URL / arXiv ID", "Notes (optional)"]):
        set_cell_bg(cell, "8FAF8F")
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"

    # Empty data rows
    for _ in range(15):
        row = table.add_row()
        for cell in row.cells:
            set_cell_bg(cell, "F7F4EF")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run("")
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(3.8)
        row.cells[1].width = Inches(2.2)

    # ── Footer note ───────────────────────────────────────────────────────────
    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(20)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Return this form to yuval.zamir@icfo.eu · incomingscience.xyz")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    run.font.italic = True
    run.font.name = "Calibri"

    out = Path("docs/incoming_science_onboarding.docx")
    out.parent.mkdir(exist_ok=True)
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
